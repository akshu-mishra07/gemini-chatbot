"""
document_processor.py - Document processing and text extraction utility module.

Supports PDF, DOCX, TXT, CSV, and image formats (PNG, JPG, JPEG, WEBP, GIF).
Extracts text and generates chunks. Falls back to OCR when standard text extraction is empty.
"""

import os
import io
import re
import sys
import csv
from PIL import Image

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "png", "jpg", "jpeg", "webp", "gif", "md", "markdown"}

def _configure_tesseract():
    """Configures pytesseract's binary path automatically on Windows."""
    import pytesseract
    import shutil
    
    # If already set or on PATH, return it
    if shutil.which("tesseract"):
        return
        
    # Common Windows installation directories for Tesseract
    common_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Tesseract-OCR\tesseract.exe")
    ]
    for path in common_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            print(f"[OCR Config] Found Tesseract at: {path}", file=sys.stderr)
            return

# Run tesseract path search on import
_configure_tesseract()


def extract_text(file_source, filename: str) -> str:
    """
    Extracts text from a file source (either a file-like object or a local file path).
    Supports PDF, DOCX, TXT, CSV, and common image formats.

    Args:
        file_source: File-like object (e.g., BytesIO) or local file path string.
        filename (str): Name of the file, used to determine the parser type.

    Returns:
        str: The full extracted text.
    """
    ext = filename.split(".")[-1].lower()
    is_path = isinstance(file_source, str)

    # 1. Read raw bytes from the file source
    if isinstance(file_source, bytes):
        file_bytes = file_source
    elif is_path:
        with open(file_source, "rb") as f:
            file_bytes = f.read()
    else:
        # File-like object (e.g. BytesIO, Streamlit UploadedFile)
        if hasattr(file_source, "getvalue"):
            file_bytes = file_source.getvalue()
        else:
            file_bytes = file_source.read()
            if hasattr(file_source, "seek"):
                file_source.seek(0)

    # 2. Extract based on extension
    if ext == "pdf":
        import pypdf
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            extracted_text = []
            for idx, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text.append(page_text)
                except Exception as page_err:
                    print(f"[PDF Extract Warning] Page {idx+1} in '{filename}' failed to extract: {page_err}", file=sys.stderr)
            full_text = "\n".join(extracted_text).strip()
            
            # Only fall back to OCR when no digital text was extracted. Short
            # text PDFs are valid and should not require Poppler/Tesseract.
            if not full_text:
                print(f"[OCR Log] PDF '{filename}' contains no extractable text. Attempting OCR...", file=sys.stderr)
                full_text = attempt_pdf_ocr(file_bytes, filename)
                
            return full_text
        except Exception as e:
            print(f"[OCR Log] Standard PDF parser failed for '{filename}': {e}. Trying OCR...", file=sys.stderr)
            try:
                return attempt_pdf_ocr(file_bytes, filename)
            except Exception as ocr_err:
                raise RuntimeError(f"Failed to extract text from PDF '{filename}' (Standard: {e}, OCR: {ocr_err})")

    elif ext in ("png", "jpg", "jpeg", "webp", "gif"):
        print(f"[OCR Log] Direct image '{filename}' uploaded. Running OCR...", file=sys.stderr)
        try:
            return attempt_image_ocr(file_bytes, filename)
        except Exception as ocr_err:
            raise RuntimeError(f"Failed to perform OCR on image '{filename}': {ocr_err}")

    elif ext == "docx":
        import docx
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX '{filename}': {e}")

    elif ext in ("txt", "md", "markdown"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from '{filename}': {e}")

    elif ext == "csv":
        try:
            decoded = file_bytes.decode("utf-8-sig", errors="ignore")
            rows = csv.reader(io.StringIO(decoded))
            lines = []
            for row_idx, row in enumerate(rows):
                cleaned = [cell.strip() for cell in row if cell and cell.strip()]
                if cleaned:
                    lines.append(f"Row {row_idx + 1}: " + " | ".join(cleaned))
            return "\n".join(lines)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from CSV '{filename}': {e}")

    else:
        raise ValueError(f"Unsupported file format: {ext}")


def attempt_pdf_ocr(file_bytes: bytes, filename: str) -> str:
    """Converts PDF pages to images and runs OCR on each page."""
    import pytesseract
    import pdf2image
    from PIL import ImageOps

    # Configure common Poppler locations for Windows
    poppler_dirs = [
        r"C:\Program Files\poppler\bin",
        r"C:\Program Files (x86)\poppler\bin",
        r"C:\poppler\bin",
        os.path.join(os.getcwd(), "poppler", "bin")
    ]
    poppler_path = None
    for d in poppler_dirs:
        if os.path.exists(d):
            poppler_path = d
            break

    try:
        images = pdf2image.convert_from_bytes(file_bytes, poppler_path=poppler_path, dpi=200)
        extracted_pages = []
        for i, page_img in enumerate(images):
            page_img = ImageOps.exif_transpose(page_img).convert("RGB")
            page_img = ImageOps.grayscale(page_img)
            page_text = pytesseract.image_to_string(page_img, config="--psm 3")
            if page_text.strip():
                extracted_pages.append(page_text)

        full_text = "\n".join(extracted_pages).strip()
        if not full_text:
            raise RuntimeError(
                f"Document '{filename}' appears to be scanned or image-only, but OCR produced no text."
            )
        return full_text
    except Exception as e:
        print(f"[OCR Error] PDF OCR failed for '{filename}': {e}", file=sys.stderr)
        raise RuntimeError(
            "OCR extraction failed. Install/configure Tesseract OCR and Poppler, "
            f"then retry '{filename}'. Reason: {e}"
        ) from e


def attempt_image_ocr(file_bytes: bytes, filename: str) -> str:
    """Runs OCR on raw image bytes."""
    import pytesseract
    from PIL import ImageOps
    try:
        img = Image.open(io.BytesIO(file_bytes))
        img = ImageOps.exif_transpose(img).convert("RGB")
        img = ImageOps.grayscale(img)
        full_text = pytesseract.image_to_string(img, config="--psm 3").strip()
        if not full_text:
            raise RuntimeError(f"Image '{filename}' was processed, but no readable text was detected.")
        return full_text
    except Exception as e:
        print(f"[OCR Error] Image OCR failed for '{filename}': {e}", file=sys.stderr)
        raise RuntimeError(
            f"OCR failed for image '{filename}'. Install/configure Tesseract OCR and retry. Reason: {e}"
        ) from e


def get_document_chunks(text: str, filename: str):
    if not text.strip():
        return []

    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_text(text)
    return [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_index": idx,
                "total_chunks": len(chunks)
            }
        )
        for idx, chunk in enumerate(chunks)
    ]
